import databutton as db
import re
from io import BytesIO
from typing import Tuple, List

from langchain.docstore.document import Document
from langchain_openai import OpenAIEmbeddings  # For embedding documents
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores.faiss import FAISS  # Vector store implementation
from pypdf import PdfReader
from docx import Document as DocxDocument
import faiss


def parse_pdf(file: BytesIO, filename: str) -> Tuple[List[str], str]:
    """
    Parses a PDF file and extracts text content page by page.
    Args:
        file (BytesIO): The PDF file as a byte stream.
        filename (str): The name of the PDF file.
    Returns:
        Tuple[List[str], str]: A tuple containing the list of extracted text pages and the filename.
    """
    pdf = PdfReader(file)
    output = []
    for page in pdf.pages:
        text = page.extract_text()
        # Cleaning and formatting the extracted text
        text = re.sub(r"(\w+)-\n(\w+)", r"\1\2", text)  # Join hyphenated words
        text = re.sub(r"(?<!\n\s)\n(?!\s\n)", " ", text.strip())  # Remove unnecessary newlines
        text = re.sub(r"\n\s*\n", "\n\n", text)  # Ensure double newlines for paragraph breaks
        output.append(text)
    return output, filename

# Function to parse docx files
def parse_docx(file: BytesIO, filename: str) -> Tuple[List[str], str]:
    """
    Parses a DOCX file and extracts text content from paragraphs.
    Args:
        file (BytesIO): The DOCX file as a byte stream.
        filename (str): The name of the DOCX file.
    Returns:
        Tuple[List[str], str]: A tuple containing the list of extracted text paragraphs and the filename.
    """
    doc = DocxDocument(file)
    output = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
    return output, filename


def text_to_docs(text: List[str], filename: str) -> List[Document]:
    """
    Converts a list of text strings into LangChain Document objects.
    Args:
        text (List[str]): A list of text strings representing pages or paragraphs.
        filename (str): The name of the file from which the text was extracted.
    Returns:
        List[Document]: A list of LangChain Document objects with metadata.
    """
    if isinstance(text, str):
        text = [text]
    page_docs = [Document(page_content=page) for page in text]
    for i, doc in enumerate(page_docs):
        doc.metadata["page"] = i + 1

    doc_chunks = []
    for doc in page_docs:
        # Splitting documents into smaller chunks for better embedding and retrieval
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=400, 
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            chunk_overlap=20,
        )
        chunks = text_splitter.split_text(doc.page_content)
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk, metadata={"page": doc.metadata["page"], "chunk": i}
            )
            doc.metadata["source"] = f"{doc.metadata['page']}-{doc.metadata['chunk']}"
            doc.metadata["filename"] = filename  # Add filename to metadata
            doc_chunks.append(doc)
    return doc_chunks


def docs_to_index(docs, openai_api_key):
    """
    Creates a FAISS vector index from a list of LangChain documents.
    Args:
        docs (List[Document]): A list of LangChain Document objects.
        openai_api_key (str): The API key for OpenAI.
    Returns:
        FAISS: The created FAISS vector index.
    """
    index = FAISS.from_documents(docs, OpenAIEmbeddings(openai_api_key=openai_api_key))
    return index


def get_index_for_pdf(pdf_and_docx_files, pdf_and_docx_names, openai_api_key):
    """
    Creates a FAISS vector index for a set of PDF and DOCX files.
    Args:
        pdf_and_docx_files (List[bytes]): A list of PDF and DOCX files as byte streams.
        pdf_and_docx_names (List[str]): A list of corresponding filenames.
        openai_api_key (str): The API key for OpenAI.
    Returns:
        FAISS: The created FAISS vector index. 
    """
    documents = []
    for file, name in zip(pdf_and_docx_files, pdf_and_docx_names):
        if name.lower().endswith('.pdf'):
            text, filename = parse_pdf(BytesIO(file), name)
        elif name.lower().endswith('.docx'):
            text, filename = parse_docx(BytesIO(file), name)
        else:
            continue  # Skip files with unsupported extensions
        documents += text_to_docs(text, filename)
    
    index = docs_to_index(documents, openai_api_key)
    return index
